import hashlib
import re
from typing import List, Dict, Any
from pathlib import Path


class DocumentParser:
    """Parse files into structured segments.

    Each segment is a dict: {"text": str, "meta": {...}} where meta carries
    structural location (page / sheet / heading) used later for citations.
    """

    @staticmethod
    def parse_file(file_path: str) -> str:
        """Backward-compatible: return plain concatenated text."""
        segments = DocumentParser.parse_segments(file_path)
        return "\n".join(s["text"] for s in segments if s["text"].strip())

    @staticmethod
    def parse_segments(file_path: str) -> List[Dict[str, Any]]:
        ext = Path(file_path).suffix.lower()

        if ext == '.txt':
            return DocumentParser._parse_txt(file_path)
        elif ext == '.md':
            return DocumentParser._parse_md(file_path)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif ext == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif ext == '.xlsx':
            return DocumentParser._parse_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _parse_txt(file_path: str) -> List[Dict[str, Any]]:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return [{"text": f.read(), "meta": {}}]

    @staticmethod
    def _parse_md(file_path: str) -> List[Dict[str, Any]]:
        """Split markdown by top-level/section headings so citations can name a section."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()

        segments: List[Dict[str, Any]] = []
        current_heading = ""
        buffer: List[str] = []

        def flush():
            text = "\n".join(buffer).strip()
            if text:
                segments.append({"text": text, "meta": {"heading": current_heading} if current_heading else {}})

        for line in content.splitlines():
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                flush()
                buffer = []
                current_heading = m.group(2).strip()
            buffer.append(line)
        flush()

        return segments or [{"text": content, "meta": {}}]

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        segments = []
        for i, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                segments.append({"text": page_text, "meta": {"page": i}})
        return segments

    @staticmethod
    def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)

        parts: List[str] = []
        current_heading = ""
        segments: List[Dict[str, Any]] = []

        def flush():
            text = "\n".join(parts).strip()
            if text:
                segments.append({"text": text, "meta": {"heading": current_heading} if current_heading else {}})

        for para in doc.paragraphs:
            if not para.text or not para.text.strip():
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                flush()
                parts.clear()
                current_heading = para.text.strip()
            parts.append(para.text)
        flush()

        # Tables: python-docx skips table text in .paragraphs, capture it explicitly
        table_lines: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
        if table_lines:
            segments.append({"text": "\n".join(table_lines), "meta": {"section": "tables"}})

        return segments

    @staticmethod
    def _parse_xlsx(file_path: str) -> List[Dict[str, Any]]:
        """Parse each sheet into row-group segments with a header line repeated.

        Instead of dumping a whole sheet as one giant blob, we emit segments of
        ~40 rows each, prefixed with the column header, so retrieval can cite a
        specific sheet + row range and chunks never cut across a row mid-cell.
        """
        import pandas as pd
        ROWS_PER_SEGMENT = 40
        dfs = pd.read_excel(file_path, sheet_name=None)
        segments: List[Dict[str, Any]] = []
        for sheet_name, df in dfs.items():
            if df.empty:
                continue
            header = " | ".join(str(c) for c in df.columns)
            total = len(df)
            for start in range(0, total, ROWS_PER_SEGMENT):
                block = df.iloc[start:start + ROWS_PER_SEGMENT]
                row_lines = [
                    " | ".join("" if pd.isna(v) else str(v) for v in row)
                    for row in block.itertuples(index=False, name=None)
                ]
                end_row = start + len(block)
                text = (
                    f"=== Sheet: {sheet_name} (hàng {start + 1}-{end_row}) ===\n"
                    f"{header}\n" + "\n".join(row_lines)
                )
                segments.append({
                    "text": text,
                    "meta": {"sheet": sheet_name, "row_start": start + 1, "row_end": end_row},
                })
        return segments


class Chunker:
    @staticmethod
    def chunk_segments(
        segments: List[Dict[str, Any]],
        chunk_size: int = None,
        overlap: int = None,
    ) -> List[Dict[str, Any]]:
        """Chunk each structured segment independently, preserving its metadata."""
        if chunk_size is None or overlap is None:
            from ..core.config import settings
            if chunk_size is None:
                chunk_size = settings.RAG_CHUNK_SIZE
            if overlap is None:
                overlap = settings.RAG_CHUNK_OVERLAP

        chunks: List[Dict[str, Any]] = []
        chunk_index = 0
        for seg in segments:
            text = seg.get("text", "")
            meta = seg.get("meta", {}) or {}
            if not text or not text.strip():
                continue
            for piece in Chunker._split(text, chunk_size, overlap):
                chunks.append({
                    "index": chunk_index,
                    "content": piece["content"],
                    "start_char": piece["start_char"],
                    "end_char": piece["end_char"],
                    "meta": meta,
                })
                chunk_index += 1
        return chunks

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """Backward-compatible plain-text chunking (no structural metadata)."""
        if not text or not text.strip():
            return []
        chunks = []
        for i, piece in enumerate(Chunker._split(text, chunk_size, overlap)):
            chunks.append({
                "index": i,
                "content": piece["content"],
                "start_char": piece["start_char"],
                "end_char": piece["end_char"],
                "meta": {},
            })
        return chunks

    @staticmethod
    def _split(text: str, chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Split text into overlapping pieces, preferring natural boundaries.

        Within each window we back off the cut point to the last paragraph
        break, sentence end, or whitespace so chunks don't start/end mid-word
        or mid-sentence (which hurts both embedding and rerank quality). Falls
        back to a hard character cut when no boundary is found.
        """
        pieces = []
        start = 0
        text_len = len(text)
        while start < text_len:
            hard_end = min(start + chunk_size, text_len)
            end = hard_end
            if hard_end < text_len:
                window = text[start:hard_end]
                # Only honor a boundary if it lands past the halfway mark, so we
                # don't produce tiny chunks.
                floor = (hard_end - start) // 2
                cut = -1
                for sep in ("\n\n", ". ", ".\n", "! ", "? ", "\n", " "):
                    idx = window.rfind(sep)
                    if idx >= floor:
                        cut = idx + len(sep)
                        break
                if cut > 0:
                    end = start + cut
            chunk = text[start:end]
            if chunk.strip():
                pieces.append({"content": chunk, "start_char": start, "end_char": end})
            if end >= text_len:
                break
            # Next window overlaps the previous one; guard against non-progress.
            next_start = end - overlap
            start = next_start if next_start > start else end
        return pieces

    @staticmethod
    def calculate_file_hash(file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
